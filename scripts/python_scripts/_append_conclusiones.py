"""_append_conclusiones.py
=========================
Añade la sección de conclusiones al final del Word existente
SIN regenerar el documento desde cero, preservando cambios manuales.

Uso:
  uv run python scripts/python_scripts/_append_conclusiones.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
PATH = ROOT / "documentation" / "Programa" / "ReporteFinal-seccionLLM.docx"

FONT = "Arial"
SIZE = 12
LINE = 276


def _set_line_spacing(p):
    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"), str(LINE))
    spacing.set(qn("w:lineRule"), "auto")


def _run(p, text, bold=False, italic=False):
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(SIZE)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return run


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_line_spacing(p)
    _run(p, text, bold=True)
    return p


def add_body_mixed(doc, parts, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_line_spacing(p)
    for text, bold, italic in parts:
        _run(p, text, bold=bold, italic=italic)
    return p


# ── Abrir documento existente ─────────────────────────────────────────────────
doc = Document(PATH)

# ── Añadir sección de conclusiones ───────────────────────────────────────────
add_heading(doc, "Conclusiones")

add_body_mixed(
    doc,
    [
        (
            "En conjunto, los tres ejes del análisis ofrecen evidencia de que el componente "
            "digital de Apapáchar genera participación genuina y no solo consumo pasivo de "
            "contenido. El 75% de los bloques temáticos codificados alcanzaron niveles básico "
            "o complejo de interacción, con la construcción moral (I6) como el indicador más "
            "activo: los temas de crianza movilizan negociación de valores de forma consistente "
            "en todos los grupos y ciudades. Las diferencias por género son coherentes con el "
            "diseño del programa — los hombres reflexionan principalmente sobre "
            "corresponsabilidad y distribución de tareas; las mujeres, sobre prácticas concretas "
            "de crianza — lo que sugiere que los contenidos activan marcos reflexivos "
            "diferenciados según el cuidador. La homogeneidad lingüística entre territorios "
            "indica, además, que el programa opera con consistencia a través de cuatro contextos "
            "urbanos distintos.",
            False,
            False,
        ),
    ],
)

add_body_mixed(
    doc,
    [
        (
            "Desde el ángulo de la escalabilidad, los hallazgos convergen en tres implicaciones "
            "operativas concretas. Primero, la deserción se concentra en las semanas 1 y 2: "
            "intervenir en esa ventana tiene el mayor retorno posible sobre la retención. "
            "Segundo, el comportamiento del facilitador es el predictor más fuerte de "
            "participación y es monitoreable en tiempo real: la rapidez de la primera "
            "respuesta de los cuidadores al inicio de cada semana anticipa con alta precisión "
            "si habrá interacción activa esa semana. Tercero, la variación en calidad de "
            "participación por módulo permite orientar la flexibilización del contenido con "
            "evidencia: los módulos de semanas 3–6 y 11–12 sostienen los niveles más altos "
            "de interacción y deberían preservarse, mientras que los de semanas 7–10 son "
            "candidatos a revisión o acortamiento. Estos tres indicadores son calculables "
            "automáticamente a partir de los metadatos de WhatsApp, sin requerir observación "
            "manual, lo que los hace viables para el monitoreo continuo a escala.",
            False,
            False,
        ),
    ],
)

# ── Guardar ───────────────────────────────────────────────────────────────────
doc.save(PATH)
print(f"Conclusiones añadidas y guardadas: {PATH}")
