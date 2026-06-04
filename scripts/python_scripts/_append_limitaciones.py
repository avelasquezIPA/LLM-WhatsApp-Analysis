"""_append_limitaciones.py
=========================
Añade frase de limitaciones al final del ultimo parrafo de conclusiones
del Word existente, SIN regenerar el documento.

Uso:
  uv run python scripts/python_scripts/_append_limitaciones.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
PATH = ROOT / "documentation" / "Programa" / "ReporteFinal-seccionLLM.docx"

FONT = "Arial"
LIMITACION = (
    " Como limitación, este es un análisis exploratorio preliminar: los marcos "
    "teóricos aplicados tienen espacio de refinamiento metodológico, y el acceso "
    "a modelos de lenguaje vía API permitiría un análisis semántico sustancialmente "
    "más potente."
)

doc = Document(PATH)

# Encontrar el último párrafo con contenido real
last_p = None
for p in reversed(doc.paragraphs):
    if p.text.strip():
        last_p = p
        break

if last_p is None:
    raise RuntimeError("No se encontró párrafo con contenido.")

print(f"Último párrafo encontrado: '{last_p.text[:80]}...'")

# Añadir la frase como un run nuevo al final del párrafo
run = last_p.add_run(LIMITACION)
run.font.name = FONT
run.font.size = Pt(12)
run.font.bold = False
run.font.italic = False
run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

doc.save(PATH)
print(f"Frase de limitaciones añadida: {PATH}")
