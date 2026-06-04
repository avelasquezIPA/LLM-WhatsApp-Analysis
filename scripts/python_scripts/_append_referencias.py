"""_append_referencias.py
========================
Añade sección de Referencias al Word existente SIN regenerarlo.

Uso:
  uv run python scripts/python_scripts/_append_referencias.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
PATH = ROOT / "documentation" / "Programa" / "ReporteFinal-seccionLLM.docx"

FONT = "Arial"
SIZE = 12
LINE = 276


def _set_line_spacing(p):
    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"), str(LINE))
    spacing.set(qn("w:lineRule"), "auto")


def _run(p, text, bold=False, italic=False):
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(SIZE)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return run


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_line_spacing(p)
    _run(p, text, bold=True)
    return p


def add_ref(doc, parts, space_after=6):
    """Parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.first_line_indent = Pt(-24)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_line_spacing(p)
    for text, bold, italic in parts:
        _run(p, text, bold=bold, italic=italic)
    return p


doc = Document(PATH)

add_heading(doc, "Referencias")

add_ref(
    doc,
    [
        (
            "Dedios-Sanguineti, M. C., Guarin, A., Torres-García, A., & Martínez Gómez, M. "
            "(2025). Assessing meaningful interaction in focus group discussions conducted over "
            "WhatsApp. ",
            False,
            False,
        ),
        ("International Journal of Qualitative Methods", False, True),
        (", 24. https://doi.org/10.1177/16094069251321599", False, False),
    ],
)

add_ref(
    doc,
    [
        (
            "Ferreira, A. A., Rocha, L., Cunha, W., et al. (2025). A comprehensive qualitative "
            "analysis of patient dialogue summarization using large language models applied to "
            "noisy, informal, non-English real-world data. ",
            False,
            False,
        ),
        ("Scientific Reports", False, True),
        (", 15, 31660. https://doi.org/10.1038/s41598-025-13560-9", False, False),
    ],
)

doc.save(PATH)
print(f"Referencias añadidas: {PATH}")
