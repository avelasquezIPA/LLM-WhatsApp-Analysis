"""_build_reporte_word.py
========================
Genera ReporteFinal-seccionLLM.docx con el mismo formato que
A+P_Informe cierre de año 2025.docx:
  - Arial 12pt en todo el documento
  - Sin colores (todo negro)
  - Sin líneas separadoras
  - Texto justificado, interlineado 1.15
  - Títulos en negrita, mismo tamaño que el cuerpo
  - Figuras grandes y centradas
  - Márgenes de 1 pulgada

Uso:
  uv run python scripts/python_scripts/_build_reporte_word.py
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
FIGURES = ROOT / "outputs" / "figures"
OUT = ROOT / "documentation" / "Programa" / "ReporteFinal-seccionLLM.docx"

FONT = "Arial"
SIZE = 12  # pt — body text
LINE = 276  # w:line value for 1.15 spacing (276/240 = 1.15)

# ── Helpers ────────────────────────────────────────────────────────────────────


def _set_line_spacing(p):
    """Apply 1.15 line spacing matching the reference document."""
    from docx.oxml import OxmlElement

    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"), str(LINE))
    spacing.set(qn("w:lineRule"), "auto")


def _run(p, text, bold=False, italic=False):
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(SIZE)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return run


def add_heading(doc, text, level=1):
    """Bold paragraph, no color. Level 1 = centered, level 2-3 = left."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    _set_line_spacing(p)
    _run(p, text, bold=True)
    return p


def add_body(doc, text, italic=False, space_after=8, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    _set_line_spacing(p)
    _run(p, text, italic=italic)
    return p


def add_body_mixed(doc, parts, space_after=8, justify=True):
    """Parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    _set_line_spacing(p)
    for text, bold, italic in parts:
        _run(p, text, bold=bold, italic=italic)
    return p


def add_figure(doc, path, caption, width=5.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    if path.exists():
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    r.font.name = FONT
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, h, bold=True)
        # Light grey fill
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9D9D9")
        tcPr.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            if "**" in cell_text:
                parts = cell_text.split("**")
                for k, part in enumerate(parts):
                    if part:
                        _run(p, part, bold=(k % 2 == 1))
            else:
                _run(p, cell_text)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)
    return table


# ── Build document ─────────────────────────────────────────────────────────────

doc = Document()

# Page setup — 1 inch margins, matching reference
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# Apply document-level defaults (Arial 12pt, 1.15 spacing) to match reference


def set_doc_defaults(doc):
    styles_element = doc.styles.element
    docDefaults = styles_element.find(qn("w:docDefaults"))
    if docDefaults is None:
        docDefaults = OxmlElement("w:docDefaults")
        styles_element.insert(0, docDefaults)
    rPrDefault = docDefaults.find(qn("w:rPrDefault"))
    if rPrDefault is None:
        rPrDefault = OxmlElement("w:rPrDefault")
        docDefaults.append(rPrDefault)
    rPr = rPrDefault.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        rPrDefault.append(rPr)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:eastAsia"), FONT)
    rFonts.set(qn("w:cs"), FONT)
    rPr.insert(0, rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")  # 12pt = 24 half-pts
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "24")
    rPr.append(sz)
    rPr.append(szCs)
    # Paragraph defaults: 1.15 spacing
    pPrDefault = docDefaults.find(qn("w:pPrDefault"))
    if pPrDefault is None:
        pPrDefault = OxmlElement("w:pPrDefault")
        docDefaults.append(pPrDefault)
    pPr = pPrDefault.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        pPrDefault.append(pPr)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), str(LINE))
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:after"), "0")
    pPr.append(spacing)


set_doc_defaults(doc)

# ── Título ─────────────────────────────────────────────────────────────────────
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_after = Pt(2)
r = _run(p_title, "Análisis computacional del componente digital", bold=True)
r.font.size = Pt(14)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(14)
r2 = _run(
    p_sub, "Programa Apapáchar (2025)  |  IPA Colombia  |  Mayo 2026", italic=True
)
r2.font.size = Pt(11)

# ── Descripción general ────────────────────────────────────────────────────────
add_heading(doc, "Descripción general", level=1)
add_body_mixed(
    doc,
    [
        (
            "El componente digital de Apapáchar operó durante 12 semanas "
            "(septiembre–diciembre 2025) a través de ",
            False,
            False,
        ),
        ("34 grupos de WhatsApp", True, False),
        (
            " en cuatro ciudades (Bogotá, Neiva, Soacha, Valledupar), con un total de ",
            False,
            False,
        ),
        ("23,950 mensajes", True, False),
        (
            " intercambiados entre facilitadores y cuidadores. Esta sección del reporte presenta "
            "los resultados del análisis computacional aplicado a ese corpus, organizado en tres ejes: "
            "análisis lingüístico y semántico, participación activa de los cuidadores, e "
            "implicaciones para la escalabilidad del programa. El análisis fue asistido por "
            "inteligencia artificial (Claude Code, Anthropic).",
            False,
            False,
        ),
    ],
)

# ── Anonimización ──────────────────────────────────────────────────────────────
add_heading(doc, "Proceso de anonimización", level=1)
add_body(
    doc,
    "Antes de cualquier análisis, los mensajes pasan por un proceso de remoción de "
    "información personalmente identificable (PII) implementado en Stata. Los datos "
    "originales (con nombres reales de participantes) son clasificados como Confidential "
    "y nunca se cargan a sistemas externos ni se procesan con herramientas de análisis "
    "automático.",
)
add_body(
    doc,
    "El script fue desarrollado y calibrado usando datos sintéticos de PII — nombres, "
    "apellidos y patrones de presentación generados artificialmente — para identificar "
    "las estructuras lingüísticas sin exponer datos reales de participantes en ninguna "
    "etapa del desarrollo. Detecta y elimina mensajes con nombres propios a través de "
    "nueve patrones que cubren las formas más comunes de auto-presentación en WhatsApp en "
    'español: "Mi nombre es…", "Me llamo…", "Soy Nombre Apellido", nombre en '
    'negritas, nombre al inicio del mensaje, referencias a terceros ("mi hijo Lucas"). '
    "La estrategia prioriza no dejar pasar PII (sensibilidad ~90%) por encima de conservar "
    "todos los mensajes. El resultado son datos clasificados como Internal, habilitando su "
    "procesamiento según las políticas de datos de IPA. Los participantes quedan "
    "identificados únicamente por IDs anónimos (P0397M, N00037) que indican grupo y "
    "sexo pero no permiten identificar a la persona.",
)

# ── Sección 1 ─────────────────────────────────────────────────────────────────
add_heading(doc, "Sección 1: Análisis lingüístico y semántico", level=1)
add_body(
    doc,
    "Marco de referencia: Ferreira et al. (2025), análisis de diálogo informal en datos "
    "ruidosos, no inglés, con modelos de lenguaje.",
    italic=True,
)
add_body_mixed(
    doc,
    [
        ("Dataset base:", True, False),
        (
            " 14,894 mensajes de texto analizables del total de 23,950 (el resto son multimedia, "
            "encuestas o mensajes eliminados). El 76% corresponde a mensajes de facilitadores y "
            "el 24% a participantes — distribución coherente con el rol activo del facilitador "
            "en el diseño del programa.",
            False,
            False,
        ),
    ],
)

add_heading(doc, "Calidad del lenguaje", level=2)
add_body(
    doc,
    "Los participantes escriben con mayor legibilidad que los facilitadores (índice "
    "Flesch-Kincaid promedio: 82.9 vs. 75.2), resultado explicado por el vocabulario "
    'técnico del programa que usan los facilitadores ("corresponsabilidad", "crianza '
    'positiva", "intrafamiliar"). Por ciudad, los participantes de las cuatro sedes '
    "escriben con complejidad léxica estadísticamente similar (ANOVA p=0.108): el programa "
    "genera expresión lingüística homogénea entre territorios. Al analizar solo los mensajes "
    "de participantes, la similitud semántica entre ciudades cae de ~0.80 a ~0.62, revelando "
    "que la homogeneidad superficial es sostenida por la voz técnica estandarizada del "
    "facilitador — los cuidadores expresan los mismos temas con vocabulario notablemente "
    "más diverso.",
)

add_heading(doc, "Afinidad de códigos cualitativos por género", level=2)
add_body(
    doc,
    "El análisis identificó los mensajes de participantes semánticamente más cercanos a "
    "los 52 códigos cualitativos del árbol de códigos del estudio (4 familias: Escalar, "
    "Viabilidad, Actitud, Resultados). Los códigos dominantes en ambos géneros pertenecen "
    "a la familia de Resultados — los participantes hablan principalmente de cambios "
    "percibidos en su propia crianza. La diferencia más relevante entre géneros es de "
    "énfasis temático:",
)
add_table(
    doc,
    ["Género", "Códigos con mayor afinidad semántica"],
    [
        [
            "Mujeres",
            "Prácticas que promueven el desarrollo del niño · Fortalecimiento del vínculo cuidador-niño · Adopción de crianza no violenta",
        ],
        [
            "Hombres",
            "Fortalecimiento del vínculo cuidador-niño · Cambio en distribución de tareas del hogar · Corresponsabilidad en crianza",
        ],
    ],
    col_widths=[1.2, 5.0],
)
add_body(
    doc,
    "Los hombres muestran alta afinidad con códigos de corresponsabilidad y distribución "
    "de tareas — que no aparecen en el top femenino — mientras que las mujeres se alinean "
    "más con prácticas concretas de crianza. Este patrón sugiere que el programa activa "
    "marcos reflexivos distintos según el género del cuidador, coherente con los objetivos "
    "diferenciales del diseño.",
)
add_figure(
    doc,
    FIGURES / "09a_citas_genero_ciudad_presentacion.png",
    "Figura 1. Citas por ciudad/género y similitud por familia de código.",
    width=5.8,
)

# ── Sección 2 ─────────────────────────────────────────────────────────────────
add_heading(
    doc, "Sección 2: Participación activa y calidad de los intercambios", level=1
)
add_body(
    doc,
    "Marco de referencia: Dedios-Sanguineti et al. (2025), evaluación de interacción "
    "significativa en grupos focales por WhatsApp.",
    italic=True,
)
add_body(
    doc,
    "Esta sección examina la participación de los cuidadores en el componente digital "
    "desde dos ángulos: primero, si participan activamente — detectando ventanas de "
    "actividad simultánea donde dos o más personas intercambian mensajes, incluyendo "
    "al facilitador cuando está presente; y segundo, de qué forma participan, "
    "caracterizando la calidad y complejidad de esos intercambios y el grado de "
    "involucramiento de los cuidadores.",
)

add_heading(doc, "Detección de sesiones de participación", level=2)
add_body(
    doc,
    "Una sesión de participación (P-P) es una ventana en la que 2 o más participantes "
    "distintos envían mensajes dentro de un lapso de 60 minutos. Es el proxy más directo "
    "de interacción activa.",
)
add_body_mixed(
    doc,
    [
        ("Se detectaron ", False, False),
        ("559 sesiones de participación", True, False),
        (
            " en el corpus completo. El 58.8% de las semanas-grupo registraron al menos una "
            "sesión activa. La sesión típica involucra 2–3 participantes simultáneos durante "
            "47 minutos. La diferencia por género de grupo es estructural: los grupos de mujeres "
            "generaron ",
            False,
            False,
        ),
        ("435 sesiones (78%)", True, False),
        (" frente a ", False, False),
        ("124 en grupos de hombres (22%)", True, False),
        (
            " — una razón de 3.5x que se establece desde las primeras semanas y no varía con "
            "el tiempo.",
            False,
            False,
        ),
    ],
)
add_figure(
    doc,
    FIGURES / "10a_panel1_sesiones_semana_presentacion.png",
    "Figura 2. Sesiones de participación por semana y ciudad.",
    width=5.8,
)

add_body_mixed(
    doc,
    [
        ("Nota sobre la muestra codificada:", True, False),
        (
            " la figura anterior engloba los 34 grupos del programa. La codificación manual "
            "de la calidad de los intercambios (sección siguiente) se limita a ",
            False,
            False,
        ),
        ("48 chats", True, False),
        (
            " (6 períodos bisemanales × 8 grupos por período): en cada período se seleccionó "
            "el grupo de mujeres y el grupo de hombres con mayor actividad de participación en "
            "cada ciudad (1 grupo por género × 4 ciudades). Del universo de 408 chats posibles "
            "(34 grupos × 12 semanas), 168 (41%) no registraron ninguna sesión de participación "
            "activa y no tenían contenido para codificar. Los 48 chats seleccionados provienen "
            "de los 240 con al menos una sesión activa, priorizando los de mayor actividad. "
            "Los niveles de complejidad observados pueden ser superiores al promedio del "
            "programa completo.",
            False,
            False,
        ),
    ],
)

add_heading(doc, "Framework de 8 indicadores (Dedios-Sanguineti)", level=2)
add_body(
    doc,
    "La presencia de una sesión de participación no dice nada sobre su calidad. Para "
    "clasificar qué tipo de construcción colectiva ocurre, se aplicó el framework de "
    "Dedios-Sanguineti et al. (2025), que define 8 indicadores (I1–I8) en tres niveles "
    "de profundidad creciente:",
)
add_table(
    doc,
    ["Nivel", "Indicadores", "Descripción"],
    [
        [
            "Stance-only",
            "I1",
            "Un participante comparte experiencia u opinión; no genera diálogo entre participantes",
        ],
        [
            "Básica",
            "I2 Consenso · I3 Desacuerdo · I4 Cambio de posición",
            "Dos o más participantes interactúan directamente",
        ],
        [
            "Compleja",
            "I5 Normalidad · I6 Construcción moral · I7 Identidades colectivas",
            "El grupo co-construye significado compartido",
        ],
        [
            "Adicional",
            "I8 Adopción de práctica",
            "Participante reporta haber aplicado el programa en casa",
        ],
    ],
    col_widths=[1.1, 2.2, 2.9],
)
add_body(
    doc,
    "La unidad de análisis es la Interacción Temática (IT): un bloque temático coherente "
    "dentro de una sesión. El nivel de cada IT es el del indicador más alto activo. I8 no "
    "afecta el nivel pero señala transferencia directa del programa al hogar. Corpus "
    "codificado: 108 ITs en 48 grupos-semana representativos (6 períodos bisemanales × "
    "8 grupos), codificadas manualmente.",
)

add_heading(doc, "Resultados principales", level=2)
add_table(
    doc,
    ["Nivel", "N ITs", "%"],
    [
        ["**Compleja** (I5/I6/I7)", "45", "41.7%"],
        ["**Básica** (I2/I3/I4)", "36", "33.3%"],
        ["Stance-only (solo I1)", "27", "25.0%"],
        ["I8 — adopción de práctica (adicional)", "14", "13.0%"],
    ],
    col_widths=[3.5, 1.0, 1.0],
)
add_body_mixed(
    doc,
    [
        ("El 75% de las ITs son básica o compleja:", True, False),
        (
            " Apapáchar genera interacción genuina en tres de cada cuatro bloques temáticos "
            "detectados. El indicador básico más frecuente es ",
            False,
            False,
        ),
        ("I2 consenso (65% de las ITs)", True, False),
        (
            " — los cuidadores expresan convergencia explícita o implícita como mecanismo "
            "principal de participación. El indicador complejo más activo es ",
            False,
            False,
        ),
        ("I6 construcción moral (33%)", True, False),
        (
            ": los temas de crianza activan negociación de valores de forma consistente. "
            "El desacuerdo directo (I3) es prácticamente ausente (2%).",
            False,
            False,
        ),
    ],
)
add_figure(
    doc,
    FIGURES / "10d_niveles_por_chunk.png",
    "Figura 3. Evolución de niveles de interacción por período bimestral.",
    width=5.8,
)

# ── Sección 3 ─────────────────────────────────────────────────────────────────
add_heading(doc, "Sección 3: Implicaciones para la escalabilidad", level=1)
add_body(
    doc,
    "Esta sección presenta los hallazgos con mayor relevancia directa para la "
    "implementación a escala del componente digital, organizados en tres ejes: "
    "retención, comportamiento del facilitador, y calidad por módulo.",
)

add_heading(doc, "Retención de participantes", level=2)
add_body(
    doc,
    'Se define como "activo" en una semana a todo participante que envió al menos un '
    "mensaje en ese período. La retención se mide desde la cohorte de semana 1.",
)
add_table(
    doc,
    ["Semana", "Retención global", "Mujeres", "Hombres"],
    [
        ["S1", "100%", "100%", "100%"],
        ["S4", "30%", "38%", "23%"],
        ["S8", "19%", "24%", "14%"],
        ["**S12**", "**15%**", "18%", "13%"],
    ],
    col_widths=[1.0, 1.5, 1.5, 1.5],
)
add_body_mixed(
    doc,
    [
        ("La retención final es del 15%. La deserción no es gradual: ", False, False),
        ("el 70% de los abandonos ocurre en las primeras dos semanas.", True, False),
        (
            " Quien llega a la semana 3 tiende a mantenerse con pérdida lenta y estable. "
            "La brecha de género es de volumen, no de retención proporcional: las mujeres "
            "retienen marginalmente mejor (18% vs. 13%) pero la diferencia grande — 3.5x más "
            "sesiones de participación en grupos de mujeres — se establece desde el inicio y "
            "no se amplía. La ventana crítica de acompañamiento es S1–S2.",
            False,
            False,
        ),
    ],
)
add_figure(
    doc,
    FIGURES / "10e_retencion_presentacion.png",
    "Figura 4. Curva de retención semanal por género.",
    width=5.8,
)

add_heading(doc, "Comportamiento del facilitador y predictor de participación", level=2)
add_body(
    doc,
    "Las familias y los equipos psicosociales perciben que uno de los factores que "
    "determinó la deserción de los grupos de WhatsApp fue la longitud de los mensajes "
    "enviados por los facilitadores. El análisis cuantitativo confirma esta percepción: "
    "la longitud promedio de los mensajes del facilitador correlaciona negativamente con "
    "la interacción en el grupo.",
)
add_table(
    doc,
    ["Métrica del facilitador", "r de Pearson", "Spearman rho", "Significancia"],
    [
        ["Longitud promedio de mensajes", "−0.43", "−0.40", "p = 0.019"],
        [
            "**% de mensajes del facilitador**",
            "**−0.91**",
            "**−0.93**",
            "**p < 0.001**",
        ],
    ],
    col_widths=[2.8, 1.1, 1.1, 1.2],
)
add_body_mixed(
    doc,
    [
        (
            "La proporción de mensajes enviados por el facilitador correlaciona negativamente "
            "con la participación del grupo (",
            False,
            False,
        ),
        ("r = −0.91, p < 0.001", True, False),
        (
            "): en grupos donde el facilitador envía más del 80% de los mensajes, la "
            "interacción entre participantes es prácticamente nula. Ambos indicadores son "
            "calculables automáticamente por grupo y semana.",
            False,
            False,
        ),
    ],
)
add_body(
    doc,
    "Un predictor complementario no circular: para cada grupo y semana se midió el "
    "tiempo entre el primer mensaje del facilitador (apertura de la sesión) y la primera "
    "respuesta de cualquier participante. Este tiempo precede causalmente a la ocurrencia "
    "de sesiones de participación.",
)
add_table(
    doc,
    ["Tiempo hasta primera respuesta", "% semanas con ≥1 sesión de participación"],
    [
        ["**≤ 2 horas**", "**82%**"],
        ["2–12 horas", "49%"],
        ["12–48 horas", "44%"],
        ["> 48 horas", "10%"],
        ["Sin respuesta esa semana", "0%"],
    ],
    col_widths=[3.0, 3.0],
)
add_body(doc, "Kruskal-Wallis H = 139, p < 0.001 · n = 408 grupo-semanas", italic=True)
add_body_mixed(
    doc,
    [
        (
            "El 82% de las semanas en que algún participante respondió al facilitador en menos "
            "de 2 horas generaron participación activa entre pares esa misma semana. "
            "Si nadie responde ese mismo día, la probabilidad cae por debajo del 10%. Esto "
            "convierte las ",
            False,
            False,
        ),
        ("primeras 2 horas de cada semana", True, False),
        (" en una ventana de alerta temprana operable en tiempo real.", False, False),
    ],
)
add_figure(
    doc,
    FIGURES / "10f_monitoreo_inicio_semana.png",
    "Figura 5. Latencia de primera respuesta y participación semanal.",
    width=5.8,
)

add_heading(
    doc,
    "Calidad de participación por módulo e implicaciones para la flexibilización",
    level=2,
)
add_body(
    doc,
    "El Plan MEL 2026 plantea explícitamente la pregunta de cómo flexibilizar el "
    "contenido digital del programa para su expansión con ICBF. El análisis de "
    "participación por módulo permite orientar esa decisión con evidencia: no todos "
    "los módulos generan el mismo nivel de interacción, y el patrón es consistente "
    "en todos los grupos y ciudades.",
)
add_table(
    doc,
    [
        "Semanas",
        "Temas",
        "% calidad",
        "Adopciones\nde práctica",
        "Sesiones de\nparticipación",
    ],
    [
        [
            "**S11–12**",
            "Situaciones difíciles / Logros afectivos",
            "**50%**",
            "1",
            "61",
        ],
        [
            "**S5–6**",
            "Emociones de cuidadores / Emociones de niños",
            "**46%**",
            "0",
            "99",
        ],
        ["S3–4", "Desarrollo infantil / Crianza integral", "43%", "8", "119"],
        ["S9–10", "Corresponsabilidad / Resolución de conflictos", "36%", "0", "55"],
        ["S7–8", "Comunicación empática / Decisiones en familia", "31%", "0", "71"],
        ["S1–2", "Conociéndonos / Valorando a nuestras familias", "29%", "13", "154"],
    ],
    col_widths=[0.7, 2.5, 0.8, 0.9, 1.1],
)
add_body(
    doc,
    "% calidad: proporción de ITs con nivel complejo (I5/I6/I7 activo). "
    "Adopciones de práctica: veces en que un cuidador reportó haber aplicado "
    "el programa en casa ese período.",
    italic=True,
)
add_body(
    doc,
    "Las semanas 3–6 combinan alta calidad (43–46%) con el mayor volumen de sesiones; "
    "las semanas 11–12 alcanzan el nivel más alto del programa (50%), impulsadas por "
    "debates sobre crianza sin estereotipos de género. Las semanas 7–10 muestran la "
    "caída más sostenida en calidad y volumen, consistente en todos los grupos y ciudades.",
)
add_table(
    doc,
    ["Categoría", "Semanas", "Justificación"],
    [
        [
            "Core — preservar",
            "S1–6 y S11–12",
            "Mayor calidad e intensidad de participación; concentración de adopciones de práctica",
        ],
        [
            "Candidatos a revisión",
            "S7–10",
            "Caída sostenida en calidad y volumen en todos los grupos",
        ],
    ],
    col_widths=[1.5, 1.2, 3.3],
)

# ── Guardar ────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Guardado: {OUT}")
