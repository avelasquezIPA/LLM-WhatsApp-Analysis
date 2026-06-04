"""_build_reporte_cualitativo.py
================================
Genera ReporteFinal-seccionCualitativa.docx en el mismo formato que
ReporteFinal-seccionLLM.docx: Arial 12pt, 1.15 interlineado, justificado,
sin colores, títulos en negrita.

Uso:
  uv run python scripts/python_scripts/_build_reporte_cualitativo.py
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "documentation" / "Programa" / "ReporteFinal-seccionCualitativa.docx"

FONT = "Arial"
SIZE = 12
LINE = 276  # 1.15 spacing


# ── Helpers ────────────────────────────────────────────────────────────────────


def _set_line_spacing(p):
    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"), str(LINE))
    spacing.set(qn("w:lineRule"), "auto")


def _run(p, text, bold=False, italic=False):
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(SIZE)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return run


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    _set_line_spacing(p)
    _run(p, text, bold=True)
    return p


def add_body(doc, text, italic=False, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_line_spacing(p)
    _run(p, text, italic=italic)
    return p


def add_body_mixed(doc, parts, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_line_spacing(p)
    for text, bold, italic in parts:
        _run(p, text, bold=bold, italic=italic)
    return p


def add_quote(doc, text, source):
    """Cita en cursiva con sangría."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_line_spacing(p)
    _run(p, f'"{text}"', italic=True)
    if source:
        _run(p, f"  —{source}", bold=False, italic=False)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, h, bold=True)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9D9D9")
        tcPr.append(shd)

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            if "**" in cell_text:
                parts = cell_text.split("**")
                for k, part in enumerate(parts):
                    if part:
                        _run(p, part, bold=(k % 2 == 1))
            else:
                _run(p, cell_text)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)
    return table


# ── Documento ─────────────────────────────────────────────────────────────────

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)


def set_doc_defaults(doc):
    styles_element = doc.styles.element
    docDefaults = styles_element.find(qn("w:docDefaults"))
    if docDefaults is None:
        docDefaults = OxmlElement("w:docDefaults")
        styles_element.insert(0, docDefaults)
    rPrDefault = docDefaults.find(qn("w:rPrDefault"))
    if rPrDefault is None:
        rPrDefault = OxmlElement("w:rPrDefault")
        docDefaults.append(rPrDefault)
    rPr = rPrDefault.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        rPrDefault.append(rPr)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:eastAsia"), FONT)
    rFonts.set(qn("w:cs"), FONT)
    rPr.insert(0, rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "24")
    rPr.append(sz)
    rPr.append(szCs)
    pPrDefault = docDefaults.find(qn("w:pPrDefault"))
    if pPrDefault is None:
        pPrDefault = OxmlElement("w:pPrDefault")
        docDefaults.append(pPrDefault)
    pPr = pPrDefault.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        pPrDefault.append(pPr)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), str(LINE))
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:after"), "0")
    pPr.append(spacing)


set_doc_defaults(doc)

# ── Título ────────────────────────────────────────────────────────────────────

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_after = Pt(2)
r = _run(
    p_title, "Resultados iniciales e intermedios — componente cualitativo", bold=True
)
r.font.size = Pt(14)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(14)
r2 = _run(
    p_sub, "Programa Apapáchar (2025)  |  IPA Colombia  |  Mayo 2026", italic=True
)
r2.font.size = Pt(11)

# ── Descripción general ───────────────────────────────────────────────────────

add_heading(doc, "Descripción general", level=1)
add_body(
    doc,
    "A lo largo de la implementación de Apapáchar, se recogió información cualitativa "
    "a través de grupos focales y entrevistas individuales con cuidadores, facilitadores "
    "y talento humano del programa. Los hallazgos son de naturaleza exploratoria y "
    "descriptiva — reflejan percepciones, actitudes y cambios auto-reportados — y "
    "complementan los resultados de la evaluación cuantitativa sin tener carácter causal. "
    "La estimación de efectos causales corresponde a la evaluación de impacto presentada "
    "en otra parte de este reporte.",
)

# ── Estrategias cualitativas ─────────────────────────────────────────────────

add_heading(doc, "Estrategias cualitativas", level=1)
add_body(
    doc,
    "La recolección de información cualitativa se llevó a cabo mediante grupos focales "
    "y entrevistas individuales con distintos actores del programa. La tabla a "
    "continuación resume las estrategias implementadas:",
    space_after=6,
)

add_table(
    doc,
    ["Estrategia", "Participantes", "N", "Detalle"],
    [
        [
            "Grupos focales",
            "Facilitadores de Fundación Apapacho",
            "2",
            "",
        ],
        [
            "Grupos focales",
            "Cuidadores participantes de Apapáchar",
            "8",
            "Bogotá: 1 (hombres)\n"
            "Soacha: 2 (1 hombres, 1 mujeres)\n"
            "Neiva: 3 (2 hombres, 1 mujeres)\n"
            "Valledupar: 2 (1 hombres, 1 mujeres)",
        ],
        [
            "Grupos focales",
            "Talento humano (TH) en Unidades de Atención",
            "4",
            "UAs donde facilitadores implementaron Apapáchar",
        ],
        [
            "Grupos focales",
            "Cuidadores participantes — Integralidad",
            "10",
            "",
        ],
        [
            "Entrevistas individuales",
            "Participantes con perfiles diferenciados",
            "5",
            "",
        ],
    ],
    col_widths=[1.5, 2.3, 0.4, 2.0],
)

# ── Hallazgos por dimensión ────────────────────────────────────────────────────

add_heading(doc, "Hallazgos por dimensión", level=1)
add_body(
    doc,
    "Los hallazgos se organizan en cuatro dimensiones temáticas que emergen de forma "
    "consistente en los grupos focales y entrevistas: gestión emocional y crianza no "
    "violenta, prácticas de crianza positiva, corresponsabilidad y equidad de género, "
    "y relaciones de pareja.",
)

# ── 1. Gestión emocional ──────────────────────────────────────────────────────

add_heading(doc, "1. Gestión emocional y crianza no violenta", level=2)
add_body(
    doc,
    "El programa fortalece la comprensión y regulación emocional de los cuidadores, "
    "favoreciendo un manejo más consciente y asertivo de sus emociones en el ejercicio "
    "de la crianza. Los participantes reportan haber adquirido estrategias concretas para "
    "pausar antes de reaccionar, lo que se traduce en una reducción auto-reportada de "
    "respuestas impulsivas. Los hombres, en particular, evidencian un cambio en sus "
    "actitudes frente al uso de prácticas de crianza violentas, reconociendo que están "
    "rompiendo patrones que les fueron transmitidos en su propia crianza.",
)
add_quote(
    doc,
    "Yo era un poquito más impulsivo, antes mis acciones eran como pegarle, hablarle "
    "fuerte, entonces con eso aprendí que es bueno a veces dejar que ellos se calmen y "
    "que ellos tomen el control.",
    "Grupo focal Hombres",
)
add_quote(
    doc,
    "Más que un cambio en mi hijo, es un cambio en mí… en la forma de actuar, enfrentar "
    "rabietas, pataletas.",
    "Grupo focal Mujeres",
)

# ── 2. Prácticas de crianza ───────────────────────────────────────────────────

add_heading(doc, "2. Prácticas de crianza positiva", level=2)
add_body(
    doc,
    "Los cuidadores reportan cambios concretos en sus prácticas cotidianas de crianza "
    "en cuatro áreas. Primero, en la interacción sensible: los hombres adoptan más "
    "espacios de juego y contacto con sus hijos e hijas, mientras que las mujeres "
    "reportan mayor empatía ante las señales de los niños. Segundo, en el ambiente de "
    "aprendizaje: los cuidadores pasan de percibir el juego y los materiales como "
    "desorden a reconocerlos como oportunidades de desarrollo. Tercero, en la disciplina: "
    "adoptan estrategias no violentas basadas en el ejemplo, la explicación y el "
    "establecimiento de rutinas. Cuarto, en el conocimiento sobre desarrollo infantil: "
    "comprenden mejor las etapas del desarrollo y manejan con mayor calma la presión "
    "social asociada a hitos del crecimiento.",
)
add_quote(
    doc,
    "Yo llegaba y mi hija tiene un montón de juguetes, y yo solamente pensaba 'otra vez "
    "haciendo desorden, qué pereza'. Entonces yo ahora llego y le pregunto, mi amor con "
    "qué juguete quiere jugar hoy, a qué vamos a jugar hoy.",
    "Grupo focal Hombres",
)
add_quote(
    doc,
    "He recordado cosas que ya había dejado de un lado… he entendido que son etapas, y "
    "que cada etapa tiene un periodo. Mi responsabilidad como cuidadora es razonar… "
    "porque mi bebé no razona.",
    "Grupo focal Mujeres",
)

# ── 3. Corresponsabilidad y género ────────────────────────────────────────────

add_heading(doc, "3. Corresponsabilidad y equidad de género", level=2)
add_body(
    doc,
    "Uno de los resultados más consistentes del programa es el cuestionamiento de "
    "estereotipos de género en la crianza. Los cuidadores reconocen que el programa les "
    "permite reflexionar sobre ideas tradicionales que limitan tanto el rol de los hombres "
    "en el cuidado como la autonomía de las mujeres para asumir tareas asociadas "
    "culturalmente a los hombres. Los participantes hombres reportan sentirse más incluidos "
    "en la crianza, asumir un rol más activo en las tareas del hogar, y resignificar los "
    "imaginarios laborales que históricamente justificaban su ausencia. Las mujeres, por "
    "su parte, reportan un cambio generacional respecto a sus propias parejas, en contraste "
    "con los patrones que observaron en sus padres.",
)
add_quote(
    doc,
    "Me toca darle comida al niño, me toca cambiar el niño, me toca bañar el niño, yo "
    "la veía a ella y yo decía no, yo también lo puedo hacer, entonces como no dejarle "
    "solamente la carga del hogar o de la casa o de los hijos a la mujer, tenemos que "
    "compartirnos la crianza porque somos los dos.",
    "Grupo focal Hombres",
)
add_quote(
    doc,
    "Mi papá no es capaz de cambiar el pañal, él se niega 'eso no es para los hombres'. "
    "En mi casa sí es distinto porque nosotros dos cambiamos a mi hija.",
    "Grupo focal Mujeres",
)
add_body(
    doc,
    "Los hombres también reportan haber cuestionado estereotipos de género transmitidos "
    "a sus hijos, reconociendo que colores, juguetes ni roles definen la identidad del "
    "niño. Las cuidadoras, por su parte, identifican su capacidad para enseñar prácticas "
    "que culturalmente se asociaban a figuras masculinas.",
)

# ── 4. Relaciones de pareja ───────────────────────────────────────────────────

add_heading(doc, "4. Relaciones de pareja y comunicación", level=2)
add_body(
    doc,
    "Los participantes evidencian un fortalecimiento en la comunicación dentro de la "
    "pareja de cuidado. El desarrollo de estrategias de autorregulación emocional en los "
    "cuidadores hombres se traduce en una mejora en el manejo de conflictos y en la "
    "apertura de espacios de diálogo más respetuosos. Los cuidadores reportan mayor "
    "empatía hacia las cargas de su pareja y una distribución más consciente de las "
    "responsabilidades del hogar.",
)
add_quote(
    doc,
    "Ahora yo respiro profundo, cuento hasta diez. Salgo, voy a respirar y me distraigo "
    "(...) La idea es dejar el tema refundido, para hablarlo más tarde de buena forma.",
    "Grupo focal Hombres",
)
add_quote(
    doc,
    "Ella me dice ¿Uy pero ese cambio qué? Pero yo le digo que no todo se lo podemos "
    "dejar a usted acá, así como yo me canso usted también se cansa, digamos como 50/50.",
    "Grupo focal Hombres",
)

# ── Nota metodológica ─────────────────────────────────────────────────────────

add_heading(doc, "Nota metodológica", level=1)
add_body(
    doc,
    "Los hallazgos presentados corresponden a resultados iniciales e intermedios "
    "recogidos durante la implementación del programa (septiembre–diciembre 2025). "
    "Al tratarse de estrategias cualitativas, los resultados reflejan la perspectiva "
    "de los participantes y no permiten establecer relaciones causales entre el programa "
    "y los cambios reportados. La consistencia de los patrones entre grupos focales "
    "realizados en cuatro ciudades y con distintos perfiles de cuidadores otorga robustez "
    "descriptiva a los hallazgos, pero su interpretación debe hacerse en conjunto con "
    "los resultados de la evaluación de impacto cuantitativa.",
)

# ── Guardar ───────────────────────────────────────────────────────────────────

doc.save(OUT)
print(f"Guardado: {OUT}")
